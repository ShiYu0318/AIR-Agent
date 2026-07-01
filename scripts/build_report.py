"""依「1142 GenAI 專題報告格式說明」產生 AIR-Agent 專題報告 DOCX。

排版規範（取自格式說明）：
- A4、上下邊距 2.54cm、左右邊距 1.91cm
- 中文標題 標楷體 16pt 粗體 置中；英文標題 Times New Roman 16pt 粗體 置中
- 作者/單位/Email 標楷體 12pt 置中
- 【摘要】【關鍵詞】 標楷體 12pt（標題粗體靠左）；摘要單段、≤500字
- 正文 標楷體 12pt、首行縮排 0.63cm、左右對齊
- 第一層標題「1. 前言」 標楷體 14pt 粗體 靠左
- 圖/表說明 標楷體 12pt 置中
- 引用文獻 第一層標題、APA 格式

用法：uv run python scripts/build_report.py
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

KAI = "標楷體"
TNR = "Times New Roman"
OUT = "專題報告_AIR-Agent.docx"


def set_cjk(run, font=KAI, size=12, bold=False, ascii_font=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = ascii_font or font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), ascii_font or font)
    rfonts.set(qn("w:hAnsi"), ascii_font or font)


def body(doc, text, indent=True, size=12, justify=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    if indent:
        pf.first_line_indent = Cm(0.63)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_cjk(p.add_run(text), size=size)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 2.0
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cjk(p.add_run(text), size=14, bold=True)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.5
    set_cjk(p.add_run(text), size=14, bold=True)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_cjk(p.add_run(text), size=12)
    return p


def bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    set_cjk(p.add_run(text), size=size)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].clear()
        set_cjk(cell.paragraphs[0].add_run(h), size=11, bold=True)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].paragraphs[0].clear()
            set_cjk(cells[i].paragraphs[0].add_run(v), size=11)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.0
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def build():
    doc = Document()

    # 版面：A4 + 邊距
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(1.91)
    sec.right_margin = Cm(1.91)

    normal = doc.styles["Normal"]
    normal.font.name = KAI
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), KAI)

    # ── 標題 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    set_cjk(p.add_run("AIR-Agent：具長期記憶與自主行為的 AI 論文研究代理人系統設計"),
            size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    set_cjk(p.add_run(
        "AIR-Agent: Designing an AI Research Agent with Long-Term Memory and Autonomous Behavior"),
        size=16, bold=True, ascii_font=TNR)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)  # 標題後空白行

    # ── 作者 ──
    for line in ["黃士育", "（單位）", "（Email）"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        set_cjk(p.add_run(line), size=12)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)  # 作者後空白行

    # ── 摘要 ──（≤500字、單段）
    abstract = (
        "本文針對自行開發的 AIR-Agent 研究代理人進行系統性深度設計解析。AIR-Agent 是一個"
        "具備自主情報蒐集、知識沉澱與報告生成能力的代理人：每日自動從 arXiv 抓取最新 AI 論文、"
        "以大型語言模型產生繁體中文摘要並推送至 Discord，同時將論文沉澱進向量庫供後續 RAG 問答；"
        "使用者亦可下達主題指令，由代理人自主檢索相關論文並產出結構化研究報告。本文涵蓋三大維度："
        "其一為架構設計，逐一拆解 System Prompt、LLM 選型、Knowledge、Plugin、Memory 與 Workflow "
        "六大支柱；其二為商業分析，以實際流程與 Token 拆解估算每日用量與運作成本，量化證明系統每月"
        "運作成本僅約 1.5 美元；其三為理論應用，探討 Token 經濟在本系統中的具體體現與設計取捨。"
        "本系統聚焦於跨語言檢索、Agentic 工作流、長期記憶沉澱與成本可控等具技術難度的議題。"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    set_cjk(p.add_run("【摘要】"), size=12, bold=True)
    set_cjk(p.add_run(abstract), size=12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_cjk(p.add_run("【關鍵詞】"), size=12, bold=True)
    set_cjk(p.add_run("研究代理人、檢索增強生成、長期記憶、Token 經濟、Discord Bot"), size=12)

    # ── 1. 前言 ──
    heading1(doc, "1. 前言")
    body(doc, "AI 領域論文以每日數百篇的速度在 arXiv 上產出，研究者面臨三個痛點：其一為資訊過載，"
              "人工追蹤新論文成本極高，且英文摘要的閱讀門檻造成資訊延遲；其二為知識零散，讀過的論文"
              "缺乏結構化沉澱，難以跨論文做關聯查詢；其三為主題調研耗時，針對新主題進行文獻回顧往往需"
              "花費數小時搜尋與閱讀。")
    body(doc, "本文將上述痛點轉化為一個可被代理人自動化的問題：給定一個持續更新的論文來源與使用者的"
              "自然語言意圖，系統應能自主決定要抓取什麼、如何濃縮、如何沉澱為可檢索知識，並在被詢問時"
              "基於可信來源生成有引用依據的回答或報告。這是一個典型的檢索增強生成（Retrieval-Augmented "
              "Generation, RAG）結合排程式自主行為（scheduled autonomy）的代理人問題，難度集中於跨語言"
              "（中文問、英文文獻）、來源可信度（須附引用）、長期記憶（避免重複收錄），以及在低成本前提下"
              "維持品質。")

    # ── 2. 系統總覽 ──
    heading1(doc, "2. 系統總覽")
    body(doc, "AIR-Agent 採分層架構：最上層為 Discord 介面層，提供 /daily、/ask、/report、/set_push_time "
              "等斜線指令；其下為 Agent 控制層，負責指令路由、排程與狀態持久化；再下為三類插件——arXiv "
              "爬蟲、LLM 推論（GroqClient）與向量庫記憶（FAISS + Embedder）；最底層為外部資源，包含 "
              "arXiv API、Groq API 與本地嵌入模型。系統以三條核心工作流串接這些元件，如表 1 所示。")

    caption(doc, "表 1. AIR-Agent 三條核心工作流")
    table(doc,
          ["工作流", "觸發方式", "行為"],
          [["每日情報推送", "排程（每日固定時間）", "抓最新論文→逐篇摘要→推送 Embed→寫入向量庫"],
           ["RAG 問答", "/ask <問題>", "向量檢索 Top-k→帶來源生成答案→附參考論文"],
           ["主題研究報告", "/report <主題>", "依主題自主檢索→多篇證據彙整→結構化報告→沉澱知識"]])

    body(doc, "技術棧方面，介面採用 discord.py 2.x（App Commands 斜線指令與 tasks.loop 排程）；推論採 Groq "
              "託管之 llama-3.3-70b-versatile（OpenAI 相容端點）；向量化採本地 sentence-transformers 的 "
              "all-MiniLM-L6-v2（384 維、離線、免費）；向量庫採 FAISS 之 IndexFlatIP（正規化後內積即餘弦"
              "相似度）；來源為 arXiv 官方 API；設定與持久化採 python-dotenv 與本地 JSON 檔。")

    # ── 3. 架構設計深度解析 ──
    heading1(doc, "3. 架構設計深度解析")
    body(doc, "本節依序拆解代理人的六大架構支柱。")

    heading2(doc, "3.1. System Prompt 設計")
    body(doc, "System Prompt 是代理人人格與行為邊界的定義。AIR-Agent 針對不同任務使用不同的 System "
              "Prompt，而非單一萬用 prompt，這是同時控制輸出品質與成本的關鍵設計。摘要任務以「明確角色、"
              "明確語言（繁體中文）、長度約束（2-3 句）、內容聚焦（問題與方法）、去客套」五項約束，既解決跨"
              "語言問題，又直接壓低輸出 token；RAG 問答任務以「僅根據提供內容回答」強制 grounding 抑制幻覺，"
              "並要求誠實退場與強制列出參考論文；主題報告任務則以 prompt 鎖定固定的 Markdown 輸出結構，使"
              "下游顯示與轉檔可預測、可解析。整體而言，System Prompt 同時扮演品質控制器與成本控制器。")

    heading2(doc, "3.2. LLM 模型選擇")
    body(doc, "模型選型權衡能力、延遲、成本、生態與開放性五個維度。摘要與多篇彙整屬中等難度、高頻率任務，"
              "70B 等級開源模型已足以勝任，旗艦閉源模型的邊際品質提升無法正當化其數十倍單價，與本系統低成本"
              "可持續運作的目標衝突。延遲是選擇 Groq 的決定性因素——其 LPU 推論延遲極低，配合 Groq 提供的 "
              "OpenAI 相容端點，可沿用既有 SDK、零學習成本。架構上 LLM 被封裝於 GroqClient，模型名稱由設定檔"
              "注入，因此模型可替換而不影響上層工作流。")

    heading2(doc, "3.3. Knowledge：知識庫佈局")
    body(doc, "知識庫採「來源→濃縮→向量化→持久化」四段佈局。其一，Chunk 策略以「一篇論文＝一個 chunk"
              "（標題＋摘要）」為粒度，因論文摘要本身已是高度濃縮的語意單元，無需再切段，避免破壞語意完整性。"
              "其二，metadata 與向量分離儲存——向量存於 FAISS、原始欄位存於 metadata.json，兩者以索引位置一一"
              "對應，檢索後回填完整 metadata 供生成與引用。其三，相似度度量採向量 L2 正規化後的內積（等於餘弦"
              "相似度），故使用 IndexFlatIP 精確檢索，在數百至數千篇的資料規模下成本可忽略且零召回損失。其四，"
              "去重機制以 arXiv 短 ID 為主鍵，新論文加入前先去重，避免知識庫膨脹。")

    heading2(doc, "3.4. Plugin：插件設計")
    body(doc, "AIR-Agent 的插件即代理人可調用的能力單元（tools），每個插件封裝一種對外部世界的動作並具清晰的"
              "輸入／輸出契約，如表 2 所示。設計上遵循單一職責與介面穩定原則：爬蟲統一回傳固定欄位的 dict，"
              "上層工作流可任意編排而無需了解內部實作；各插件皆具可替換性，爬蟲可擴充其他來源、LLM 可換供應商、"
              "向量庫可換 Pinecone 或 Qdrant，皆不影響工作流契約。")

    caption(doc, "表 2. AIR-Agent 插件（能力單元）一覽")
    table(doc,
          ["插件", "封裝類別", "能力", "對外資源"],
          [["arXiv 爬蟲", "ArxivCrawler", "fetch_latest_papers / search_topic", "arXiv API"],
           ["LLM 推論", "GroqClient", "summarize / answer / research_report", "Groq API"],
           ["向量庫", "VectorStore", "add 收錄、search 檢索", "本地 FAISS"],
           ["嵌入器", "Embedder", "encode 文本向量化", "本地模型"],
           ["排程器", "tasks.loop", "定時觸發每日推送", "系統時鐘"],
           ["指令介面", "App Commands", "/daily、/ask、/report、/set_push_time、/help", "Discord"]])

    heading2(doc, "3.5. Memory：記憶機制")
    body(doc, "記憶是區分無狀態 chatbot 與代理人的核心。AIR-Agent 設計三層記憶：長期語意記憶（FAISS 向量庫＋"
              "metadata.json，永久存活，沉澱讀過的論文以支撐 RAG）、狀態記憶（schedule.json 與去重集合，記住"
              "使用者設定的推送時間與已收錄論文）、工作記憶（單次請求的 context window，把檢索到的 Top-k 證據塞"
              "入 prompt）。關鍵在於長期記憶以本地向量庫實現，意味著「記住」不消耗 API token，只有在「回憶並使用」"
              "（檢索後塞入 prompt）時才產生成本——這是本系統 Token 經濟的核心優勢。")

    heading2(doc, "3.6. Workflow：工作流控制")
    body(doc, "三條工作流的工程重點包含：非阻塞執行（耗時的爬蟲與 LLM 呼叫透過 asyncio.to_thread 丟至背景執行緒，"
              "避免阻塞 Discord 事件迴圈）、逾時防護（互動指令一律先 defer()，把 Discord 的 3 秒硬限制延展為 15 "
              "分鐘的 followup 視窗）、容錯不中斷（排程找不到頻道或 LLM 呼叫失敗皆有 try/except 與降級回覆，單點"
              "失敗不致服務崩潰）、動態重排程（/set_push_time 透過 change_interval() 即時改排程，無需重啟）。")

    # ── 4. 安全性與工程取捨 ──
    heading1(doc, "4. 安全性與工程取捨")
    body(doc, "本系統在多個面向上做了明確的工程取捨，如表 3 所示。整體取向是在學術可信、低成本與工程簡潔之間求"
              "取平衡。")
    caption(doc, "表 3. 安全性與工程取捨")
    table(doc,
          ["議題", "取捨", "說明"],
          [["幻覺抑制", "grounding＋強制引用", "以 prompt 約束僅根據來源並列出參考，學術可信"],
           ["金鑰管理", ".env 不進版控", "Groq／Discord token 透過環境變數注入"],
           ["精確 vs 近似檢索", "選精確（FlatIP）", "資料規模小，零召回損失優先於檢索速度"],
           ["雲端 vs 本地嵌入", "選本地", "免費、離線、無 API token 成本"],
           ["單一 vs 多 System Prompt", "選多", "各任務專屬 prompt，品質與成本雙控"]])

    # ── 5. 商業分析 ──
    heading1(doc, "5. 商業分析：Token 用量與運作成本估算")
    body(doc, "本節單價以 Groq 對 llama-3.3-70b-versatile 的公開參考定價計算（輸入約 0.59 美元／百萬 token、"
              "輸出約 0.79 美元／百萬 token）。以實際 prompt 結構與 arXiv 摘要的典型長度估算單次操作的 token 用量，"
              "如表 4 所示。")
    caption(doc, "表 4. 單次操作的 Token 拆解（估）")
    table(doc,
          ["操作", "輸入 token", "輸出 token", "組成說明"],
          [["summarize（單篇）", "≈ 380", "≈ 150", "system 80＋標題／摘要 300；輸出 2-3 句中文"],
           ["answer（單次問答）", "≈ 1,430", "≈ 400", "system 120＋4 篇證據 1,280＋問題 30"],
           ["research_report（單次）", "≈ 3,070", "≈ 1,200", "system 250＋8 篇證據 2,800＋主題 20"]])
    body(doc, "假設一個中小型 AI 研究社群的日常使用：每日排程推送 1 次（5 篇摘要）、/ask 問答 30 次、/report 報告 "
              "5 次。據此估算每日總輸入約 60,150 token、輸出約 18,750 token，合計約 78,900 token（約 0.079 M）。"
              "對應每日輸入成本約 0.0355 美元、輸出成本約 0.0148 美元，每日總成本約 0.050 美元（約新台幣 1.6 元），"
              "每月總成本約 1.51 美元（約新台幣 48 元）。向量化使用本地模型，不經 API、不計 token 費用。")
    body(doc, "成本結構洞察有三：其一，問答（/ask）是主要成本來源，佔輸入 token 約 71%，因每次須將 4 篇證據塞入 "
              "context；其二，輸出 token 雖單價較高但佔比僅約 24%，歸功於 System Prompt 的長度約束；其三，成本與"
              "「檢索證據數 k」及「報告論文數」近似線性相關，是最直接的成本旋鈕。在 10 倍與 100 倍使用量情境下，"
              "每月成本分別約為 15 美元與 150 美元，呈線性擴展。")

    # ── 6. 理論應用 ──
    heading1(doc, "6. 理論應用：Token 經濟（Tokenomics）")
    body(doc, "本文所指 Token 經濟，是將 token 視為 LLM 系統中的基本計價與資源單位，研究如何在 token 的供給"
              "（context window 容量）、需求（任務所需資訊）與價格（輸入／輸出單價）之間做最佳化配置，以在品質、"
              "延遲、成本三者間取得最優解。AIR-Agent 的設計通篇貫穿此思維，具體體現於五個核心概念。")
    body(doc, "第一，輸入／輸出的非對稱定價：LLM 普遍對輸出 token 收取高於輸入的單價，因自迴歸生成中每個輸出 "
              "token 都需一次完整前向傳播。本系統透過 System Prompt 強制壓縮輸出，把成本壓在較便宜的輸入側。")
    body(doc, "第二，Context Window 是稀缺資源：容量有限且越長越貴、越長越慢、越易迷失。RAG 即 Token 經濟的核心"
              "手段——用便宜的本地向量檢索篩出 Top-k 最相關證據，只把這幾百至幾千 token 餵給 LLM，而非整個知識庫，"
              "這是「用檢索成本換取生成成本」的經典套利。")
    body(doc, "第三，記憶外部化：把長期知識存在向量庫而非塞進 prompt，等於把記憶從按 token 計價的昂貴 context 搬"
              "到一次性計算、永久免費複用的本地儲存，記憶的邊際持有成本趨近於零。第四，邊際成本與規模經濟：LLM "
              "API 邊際成本隨用量近似線性，故單位請求的 token 效率直接決定可規模化的天花板。第五，模型分層與成本／"
              "能力配適：選用 70B 開源模型而非旗艦閉源模型，在中等難度任務上取得最佳能力／成本比。")
    body(doc, "進階機制方面，本系統可延伸採用 Prompt Caching（快取固定的 system 前綴）、語意快取（相似問題直接回傳"
              "快取答案）、輸出長度上限（防止失控生成）與批次嵌入（攤平固定開銷）等，進一步壓低成本。")

    # ── 7. 限制與未來工作 ──
    heading1(doc, "7. 限制與未來工作")
    body(doc, "現有限制包含：來源單一（僅 arXiv），主題覆蓋受限；檢索為純向量（dense）檢索，對專有名詞與縮寫的"
              "精確匹配較弱；摘要與報告品質受 70B 模型上限約束。未來工作方向如下：")
    bullet(doc, "多來源整合：HuggingFace Papers、會議論文、GitHub Trending。")
    bullet(doc, "混合檢索（Hybrid Search）：dense 結合 BM25 sparse，提升專有名詞召回。")
    bullet(doc, "Re-ranking：檢索後加入 cross-encoder 重排，提升 Top-k 精準度。")
    bullet(doc, "多語言嵌入升級：針對中文查詢與英文文獻的跨語言檢索優化。")
    bullet(doc, "Prompt 與語意快取：落實成本優化機制。")
    bullet(doc, "趨勢分析：對沉澱的論文做時間序列主題演化分析。")

    # ── 8. 結論 ──
    heading1(doc, "8. 結論")
    body(doc, "AIR-Agent 以「自動情報蒐集→知識沉澱→RAG 問答→主題報告」的完整閉環，展示了一個具備長期記憶與"
              "自主行為的研究代理人應有的架構深度。在架構面，六大支柱各司其職且鬆耦合、可替換；在商業面，透過"
              "嚴謹的 token 拆解證明系統可在每月約 1.5 美元的極低成本下持續運作；在理論面，系統設計處處體現 "
              "Token 經濟的核心思想——以檢索換生成、以本地化歸零記憶成本、以 prompt 工程壓縮輸出。本專案不僅是"
              "一個可運作的系統，更是一個關於如何在成本約束下設計可持續 AI 代理人的具體論證。")

    # ── 引用文獻 ──
    heading1(doc, "引用文獻")
    refs = [
        "Grattafiori, A., et al. (2024). The Llama 3 herd of models. arXiv:2407.21783.",
        "Johnson, J., Douze, M., & Jégou, H. (2019). Billion-scale similarity search with GPUs. "
        "IEEE Transactions on Big Data, 7(3), 535-547.",
        "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. "
        "Advances in Neural Information Processing Systems, 33, 9459-9474.",
        "Liu, N. F., et al. (2024). Lost in the middle: How language models use long contexts. "
        "Transactions of the Association for Computational Linguistics, 12, 157-173.",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese "
        "BERT-networks. In Proceedings of EMNLP-IJCNLP (pp. 3982-3992).",
    ]
    for r in refs:
        p = doc.add_paragraph(style="Quote")
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(4)
        set_cjk(p.add_run(r), size=11, ascii_font=TNR)

    doc.save(OUT)
    print(f"已輸出：{OUT}")


if __name__ == "__main__":
    build()
