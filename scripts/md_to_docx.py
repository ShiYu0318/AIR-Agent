"""輕量 Markdown → DOCX 轉換器（針對 report.md 使用的語法）。

支援：# ## ### 標題、段落、- 與 1. 清單、| 表格 |、```程式碼區塊```、
水平線 ---、行內 **粗體** 與 `行內碼`。

用法：uv run python scripts/md_to_docx.py report.md report.docx
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text):
    """處理 **粗體** 與 `行內碼`。"""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(part)


def is_table_sep(line):
    return bool(re.fullmatch(r"\s*\|?[\s:|-]+\|?\s*", line)) and "-" in line


def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def add_table(doc, rows):
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]  # 略過分隔線
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        add_inline(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            text = r[i] if i < len(r) else ""
            p = cells[i].paragraphs[0]
            add_inline(p, text)
    doc.add_paragraph()


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 程式碼區塊
        if stripped.startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        # 表格：當前行像表格列，且下一行是分隔線
        if "|" in line and i + 1 < n and is_table_sep(lines[i + 1]):
            rows = [line, lines[i + 1]]
            i += 2
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            continue

        # 標題
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1:
                h = doc.add_heading("", level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(h, text)
            else:
                h = doc.add_heading("", level=min(level - 1, 4))
                add_inline(h, text)
            i += 1
            continue

        # 水平線
        if re.fullmatch(r"-{3,}", stripped):
            doc.add_paragraph().add_run("─" * 40)
            i += 1
            continue

        # 無序清單
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(2))
            i += 1
            continue

        # 有序清單
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run_text = stripped.lstrip(">").strip()
            add_inline(p, run_text)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 一般段落
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"已輸出：{docx_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "report.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else "report.docx"
    convert(src, dst)
